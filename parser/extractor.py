"""Extract raw text from uploaded resume files (PDF or Word)."""

import os

import pdfplumber
import docx


def extract_text(file_path: str) -> str:
    """Extract raw text from a PDF or Word document.

    Raises ValueError for unsupported file types.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_from_pdf(file_path: str) -> str:
    text_chunks = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_chunks.append(page_text)
    return "\n".join(text_chunks)


def _extract_from_docx(file_path: str) -> str:
    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs]

    # Many resumes use tables for layout (e.g. skills/contact columns),
    # so pull that text out too.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)
